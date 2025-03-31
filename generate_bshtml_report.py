from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

def generate_bshtml_analysis_report(output_path="d:\\project\\data_clean\\docs\\BSHTMLLoader问题分析报告.docx"):
    """生成BSHTMLLoader问题分析的Word报告"""
    
    # 创建文档对象
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 添加标题
    title = doc.add_heading('BSHTMLLoader问题分析报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加报告生成时间
    time_paragraph = doc.add_paragraph()
    time_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    time_run = time_paragraph.add_run(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # 添加问题概述
    doc.add_heading('1. 问题概述', level=1)
    doc.add_paragraph('在使用langchain_community.document_loaders中的BSHTMLLoader处理HTML文件时，持续出现以下问题：')
    
    problem_list = doc.add_paragraph(style='List Bullet')
    problem_list.add_run('文件编码错误 (UnicodeDecodeError)')
    doc.add_paragraph(style='List Bullet').add_run('文件路径不存在错误（当处理内存内容时）')
    doc.add_paragraph(style='List Bullet').add_run('元数据信息丢失或不准确')
    
    # 添加根本原因分析
    doc.add_heading('2. 根本原因分析', level=1)
    
    # 编码问题
    doc.add_heading('2.1 编码处理缺陷', level=2)
    doc.add_paragraph('BSHTMLLoader在实现中存在编码处理缺陷：')
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('''# BSHTMLLoader源码片段（简化）
def __init__(self, file_path: str):
    with open(file_path, "r") as f:  # 未指定编码参数
        soup = BeautifulSoup(f, "html.parser")''')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    doc.add_paragraph('关键问题：')
    doc.add_paragraph(style='List Bullet').add_run('依赖系统默认编码（Windows默认使用GBK/cp936）')
    doc.add_paragraph(style='List Bullet').add_run('未实现编码自动检测机制')
    doc.add_paragraph(style='List Bullet').add_run('没有编码回退策略')
    
    # 路径处理问题
    doc.add_heading('2.2 路径处理问题', level=2)
    doc.add_paragraph('BSHTMLLoader在元数据处理上存在设计缺陷：')
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('''# 元数据记录实现（简化）
self.metadata = {"source": file_path}  # 强制依赖物理文件路径''')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    doc.add_paragraph('导致后果：')
    doc.add_paragraph(style='List Bullet').add_run('处理内存内容时丢失原始路径信息')
    doc.add_paragraph(style='List Bullet').add_run('临时文件路径污染元数据')
    doc.add_paragraph(style='List Bullet').add_run('无法直接处理字符串内容')
    
    # 解决方案尝试
    doc.add_heading('3. 解决方案尝试', level=1)
    
    # 方案1
    doc.add_heading('3.1 尝试1：指定编码参数', level=2)
    doc.add_paragraph('首先尝试通过指定编码参数解决问题：')
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('''# 尝试指定编码
with open(file_path, 'r', encoding='utf-8') as f:
    html_content = f.read()
loader = BSHTMLLoader(html_content)  # 错误：期望文件路径而非内容''')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    doc.add_paragraph('结果：失败。BSHTMLLoader期望接收文件路径而非内容。')
    
    # 方案2
    doc.add_heading('3.2 尝试2：编码回退机制', level=2)
    doc.add_paragraph('尝试实现编码回退机制：')
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('''def _read_file_with_encoding(self, file_path: str) -> str:
    """尝试不同编码读取文件"""
    encodings = ['utf-8', 'gbk', 'latin-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue''')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    doc.add_paragraph('结果：部分成功，但BSHTMLLoader内部仍然使用默认编码重新读取文件。')
    
    # 方案3
    doc.add_heading('3.3 尝试3：临时文件桥接', level=2)
    doc.add_paragraph('尝试使用临时文件作为桥接：')
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('''# 创建临时文件
with tempfile.NamedTemporaryFile(suffix='.html', mode='w', encoding='utf-8', delete=False) as temp:
    temp.write(html_content)
    temp_path = temp.name

# 使用临时文件创建BSHTMLLoader
loader = BSHTMLLoader(temp_path)
documents = loader.load()''')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    doc.add_paragraph('结果：仍然存在问题，BSHTMLLoader内部可能使用不同的编码重新读取文件。')
    
    # 最终解决方案
    doc.add_heading('4. 推荐解决方案', level=1)
    doc.add_paragraph('经过多次尝试，最有效的解决方案是放弃使用BSHTMLLoader，直接使用BeautifulSoup处理HTML内容：')
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('''from bs4 import BeautifulSoup
from langchain_core.documents import Document

def parse_html(html_content, file_path):
    """解析HTML内容并返回Document对象"""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 获取正文内容
    text = soup.get_text(separator='\\n', strip=True)
    
    # 提取元数据
    metadata = {
        'source': file_path,
        'title': soup.title.string if soup.title else '',
        'language': soup.get('lang', ''),
    }
    
    return Document(
        page_content=text,
        metadata=metadata
    )''')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    # 解决方案对比
    doc.add_heading('5. 解决方案对比', level=1)
    
    # 创建表格
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # 添加表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '方案'
    hdr_cells[1].text = '优点'
    hdr_cells[2].text = '缺点'
    
    # 添加方案1
    row_cells = table.add_row().cells
    row_cells[0].text = '使用临时文件桥接'
    row_cells[1].text = '保持与LangChain生态兼容'
    row_cells[2].text = '增加IO操作，存在路径泄露风险'
    
    # 添加方案2
    row_cells = table.add_row().cells
    row_cells[0].text = '直接使用BeautifulSoup'
    row_cells[1].text = '完全控制解析流程，高效'
    row_cells[2].text = '需要重新实现元数据逻辑'
    
    # 添加方案3
    row_cells = table.add_row().cells
    row_cells[0].text = '修改Loader源码'
    row_cells[1].text = '根本解决问题'
    row_cells[2].text = '需要维护自定义版本'
    
    # 结论
    doc.add_heading('6. 结论', level=1)
    conclusion = doc.add_paragraph()
    conclusion.add_run('BSHTMLLoader存在设计缺陷，主要体现在编码处理和路径处理上。最佳解决方案是直接使用BeautifulSoup处理HTML内容，并手动创建Document对象。这种方法可以完全控制解析流程，避免BSHTMLLoader内部实现的限制。')
    
    # 确保目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # 保存文档
    doc.save(output_path)
    print(f"报告已生成：{output_path}")

if __name__ == "__main__":
    generate_bshtml_analysis_report()
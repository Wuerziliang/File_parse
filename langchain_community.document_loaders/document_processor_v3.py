"""
文档处理器模块，支持多种文件格式的加载和处理
包括：PDF、TXT、CSV、DOC/DOCX、JSON、图片(PNG/JPG/JPEG)以及压缩文件(ZIP/7Z)
"""

# 标准库导入
import os
import json
import zipfile
import tempfile
import py7zr
import base64
import threading
import _thread
from datetime import datetime
from typing import Dict, Any, List, Optional

# 第三方库导入
from tqdm import tqdm
from docx import Document as DocxDocument
import dashscope
from dashscope import MultiModalConversation

# LangChain相关导入
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    CSVLoader, PyPDFLoader, TextLoader, 
    JSONLoader, Docx2txtLoader
)

class DocumentProcessor:
    """文档处理主类，负责加载和处理各种格式的文档"""
    
    def __init__(self, input_dir: str, output_dir: str):
        """
        初始化文档处理器
        
        参数:
            input_dir: 输入目录路径
            output_dir: 输出目录路径
        """
        # 初始化DashScope API密钥
        dashscope.api_key = os.getenv('DASHSCOPE_API_KEY')
        if not dashscope.api_key:
            raise ValueError("DASHSCOPE_API_KEY环境变量未设置")
            
        self.input_dir = input_dir
        self.output_dir = output_dir
        
        # 支持的文件格式及对应的处理器
        self.supported_formats = {
            'pdf': PyPDFLoader,
            'txt': TextLoader,
            'csv': CSVLoader,
            'doc': self._process_doc,
            'docx': Docx2txtLoader,
            'json': lambda p: JSONLoader(p, jq_schema='.', text_content=False),
            'png': self._process_image,
            'jpg': self._process_image,
            'jpeg': self._process_image
        }
        
        # 创建输出目录结构
        for fmt in (*self.supported_formats.keys(), 'zip', '7z'):
            os.makedirs(os.path.join(output_dir, fmt), exist_ok=True)

    def _process_doc(self, file_path: str) -> List[Document]:
        """处理DOC/DOCX文档"""
        if file_path.endswith('.docx'):
            # 使用python-docx处理DOCX文件
            doc = DocxDocument(file_path)
            return [Document(
                page_content='\n'.join(p.text for p in doc.paragraphs),
                metadata={"source": file_path}
            )]
        # 使用TextLoader处理旧版DOC文件
        return TextLoader(file_path).load()

    def _process_image(self, file_path: str) -> List[Document]:
        """使用DashScope API处理图片文件，提取文字内容"""
        try:
            # 读取并编码图片
            with open(file_path, 'rb') as f:
                encoded_image = base64.b64encode(f.read()).decode('utf-8')
            
            # 调用OCR API
            response = MultiModalConversation.call(
                model='qwen-vl-ocr',
                messages=[{
                    'role': 'user',
                    'content': [{'image': f'data:image/jpeg;base64,{encoded_image}'}]
                }]
            )
            
            if response.status_code == 200:
                # 处理API响应
                content = response.output.choices[0].message.content
                text_content = '\n'.join(
                    item['text'] for item in content if isinstance(content, list) and 'text' in item
                ) if isinstance(content, list) else str(content)
                
                return [Document(
                    page_content=text_content,
                    metadata={"source": file_path, "processor": "qwen-vl-ocr"}
                )]
            
            print(f"OCR失败: {response.status_code} - {response.message}")
        except Exception as e:
            print(f"图片处理错误: {e}")
        return []

    def process_file(self, file_path: str) -> Optional[Dict[str, Any]]:
        """处理单个文件，返回处理结果或None"""
        file_ext = file_path.split('.')[-1].lower()
        
        # 处理压缩文件
        if file_ext in ['zip', '7z']:
            return self._process_archive(file_path, file_ext)
            
        # 检查是否支持该格式
        if file_ext not in self.supported_formats:
            return None

        # 设置处理超时(120秒)
        timer = threading.Timer(120, _thread.interrupt_main)
        timer.start()
        try:
            # 根据文件类型选择处理器
            if file_ext in ['png', 'jpg', 'jpeg']:
                content = self._process_image(file_path)
            else:
                loader = self.supported_formats[file_ext](file_path)
                content = loader.load()
            
            return {
                'content': content,
                'file_type': file_ext,
                'original_path': file_path
            }
        except Exception as e:
            print(f"处理文件 {file_path} 失败: {type(e).__name__} - {str(e)}")
            return None
        finally:
            timer.cancel()

    def _process_archive(self, file_path: str, file_ext: str) -> Optional[Dict[str, Any]]:
        """处理压缩文件(ZIP/7Z)"""
        with tempfile.TemporaryDirectory() as temp_dir:
            # 选择解压工具
            extractor = zipfile.ZipFile if file_ext == 'zip' else py7zr.SevenZipFile
            
            # 解压文件
            with extractor(file_path, 'r') as archive:
                archive.extractall(temp_dir)
            
            # 处理压缩包内文件
            content = []
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_result = self.process_file(os.path.join(root, file))
                    if file_result:
                        content.append(file_result)
                        # 保存内部文件处理结果
                        self.save_result({
                            'content': file_result['content'],
                            'file_type': file_ext,
                            'original_path': os.path.join(file_path, file)
                        })
            
            # 保存压缩包整体处理结果
            if content:
                result = {
                    'content': content,
                    'file_type': file_ext,
                    'original_path': file_path
                }
                self.save_result(result)
                return result
        return None

    def save_result(self, result: Dict[str, Any]) -> bool:
        """保存处理结果到JSON文件"""
        if not result:
            return False
            
        try:
            # 生成输出路径
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(
                self.output_dir, 
                result['file_type'],
                f"{os.path.basename(result['original_path'])}_{timestamp}.json"
            )
            
            # 统一处理不同格式的文档内容
            documents = []
            for doc in result['content']:
                if isinstance(doc, dict):  # 压缩文件内容
                    doc_entry = {
                        'content': str(doc.get('content', '')),
                        'metadata': doc.get('metadata', {})
                    }
                else:  # Document对象
                    doc_entry = {
                        'content': doc.page_content,
                        'metadata': doc.metadata
                    }
                documents.append(doc_entry)
            
            # 写入JSON文件
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump({
                    'original_file': result['original_path'],
                    'processed_time': timestamp,
                    'documents': documents
                }, f, ensure_ascii=False, indent=2)
            return True
            
        except Exception as e:
            print(f"保存结果时出错: {str(e)}")
            return False

    def process_directory(self):
        """处理整个目录下的文件"""
        # 收集支持和不支持的文件
        supported_files = []
        skipped_files = []
        
        for root, _, files in os.walk(self.input_dir):
            for file in files:
                ext = file.split('.')[-1].lower()
                if ext in {**self.supported_formats, 'zip': None, '7z': None}:
                    supported_files.append((root, file))
                else:
                    skipped_files.append(os.path.join(root, file))
        
        # 打印跳过的文件
        if skipped_files:
            print("\n跳过的文件：")
            for file in skipped_files:
                print(f"- {file}")
        
        # 处理文件并显示进度
        total_files = sum(1 for _ in os.walk(self.input_dir) for _ in _[2])
        success_count = failed_count = 0
        
        with tqdm(total=len(supported_files), desc="处理文件") as pbar:
            for root, file in supported_files:
                file_path = os.path.join(root, file)
                pbar.set_postfix_str(f"正在处理: {file}")
                
                if result := self.process_file(file_path):
                    # 处理压缩文件内容
                    if isinstance(result['content'], list) and all(isinstance(x, dict) for x in result['content']):
                        success_count += sum(bool(self.save_result(zip_result)) for zip_result in result['content'])
                    else:
                        success_count += bool(self.save_result(result))
                else:
                    failed_count += 1
                pbar.update(1)
        
        # 输出处理结果统计
        print(f"\n处理完成：")
        print(f"总文件数：{total_files}")
        print(f"成功：{success_count}")
        print(f"失败：{failed_count}")
        print(f"跳过：{total_files - len(supported_files)}")

def main():
    """主函数"""
    processor = DocumentProcessor(
        "d:\\project\\data_clean\\input",
        "d:\\project\\data_clean\\output"
    )
    processor.process_directory()

if __name__ == "__main__":
    main()
    
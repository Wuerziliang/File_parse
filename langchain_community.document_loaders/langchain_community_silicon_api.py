from langchain_community.document_loaders import (
    CSVLoader, TextLoader, 
    JSONLoader, Docx2txtLoader
)
import os, json, zipfile, tempfile
import py7zr, textract
from datetime import datetime
from typing import Dict, Any, List
import threading, _thread
from tqdm import tqdm
from langchain_core.documents import Document
import requests
import fitz  # PyMuPDF
import base64
from io import BytesIO
from PIL import Image
from docx import Document as DocxDocument

class QwenPDFLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.api_url = "https://api.siliconflow.cn/v1/chat/completions"
        self.headers = {
            "Authorization": "Bearer sk-jelifbfcfupdddnoqzgunzehavucvjzcpsmqmwoirqbacehl",
            "Content-Type": "application/json"
        }
        self.api_params = {
            "model": "Qwen/QwQ-32B",
            "stream": False,
            "max_tokens": 512,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "frequency_penalty": 0.5,
            "n": 1
        }
        
    def _extract_images_from_page(self, page) -> List[str]:
        image_list = []
        for image_index, img in enumerate(page.get_images()):
            xref = img[0]
            base_image = page.parent.extract_image(xref)
            image_bytes = base_image["image"]
            
            image = Image.open(BytesIO(image_bytes))
            buffered = BytesIO()
            image.save(buffered, format="PNG")
            img_str = base64.b64encode(buffered.getvalue()).decode()
            image_list.append(img_str)
        return image_list

    def load(self) -> List[Document]:
        doc = fitz.open(self.file_path)
        documents = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            images = self._extract_images_from_page(page)
            
            payload = {
                **self.api_params,
                "messages": [{
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": f"请提取并总结这个PDF页面的主要内容：\n{text}"
                        }
                    ] + [
                        {
                            "type": "image",
                            "image": img_base64
                        } for img_base64 in images
                    ]
                }]
            }
            
            try:
                response = requests.post(
                    self.api_url,
                    headers=self.headers,
                    json=payload
                )
                
                if response.status_code == 200:
                    result = response.json()
                    processed_content = result["choices"][0]["message"]["content"]
                    documents.append(Document(
                        page_content=processed_content,
                        metadata={
                            "source": self.file_path,
                            "page": page_num + 1
                        }
                    ))
            except Exception as e:
                print(f"处理页面 {page_num + 1} 时出错: {str(e)}")
            
        doc.close()
        return documents

class DocumentProcessor:
    def __init__(self, input_dir: str, output_dir: str):
        self.input_dir = input_dir
        self.output_dir = output_dir
        self.supported_formats = {
            'pdf': QwenPDFLoader,  # 替换为新的QwenPDFLoader
            'txt': TextLoader, 
            'csv': CSVLoader,
            'doc': lambda path: self._process_doc(path),
            'docx': Docx2txtLoader,
            'json': lambda path: JSONLoader(path, jq_schema='.', text_content=False)
        }
        [os.makedirs(os.path.join(output_dir, fmt), exist_ok=True) 
         for fmt in {**self.supported_formats, 'zip': None, '7z': None}]

    def _process_doc(self, file_path: str):
        if file_path.endswith('.docx'):
            doc = DocxDocument(file_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        else:  # 对于 .doc 文件仍然使用 textract
            text = textract.process(file_path).decode('utf-8')
        return [Document(page_content=text, metadata={"source": file_path})]